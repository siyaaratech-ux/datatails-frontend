from flask import Flask, request, jsonify
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import json
import traceback
import os
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import logging
import secrets

# Load environment variables from multiple locations (later overrides earlier)
_load_dir = os.path.dirname(os.path.abspath(__file__))
for path in [
    os.path.join(_load_dir, '..', '.env'),
    os.path.join(_load_dir, '.env'),
    os.path.join(os.getcwd(), '.env'),
]:
    if os.path.isfile(path):
        load_dotenv(path)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Lazy imports for heavy libraries
_kg_chat = None
_vrs = None
_PyPDF2 = None
_Document = None
_pd = None
_groq = None
_CronTab = None
_np = None

def import_kg_chat():
    global _kg_chat
    if _kg_chat is None:
        import kg_chat
        _kg_chat = kg_chat
    return _kg_chat

def import_vrs():
    global _vrs
    if _vrs is None:
        import vrs
        _vrs = vrs
    return _vrs

def get_groq_client():
    global _groq
    if _groq is None:
        import groq
        _groq = groq
    global client
    if 'client' not in globals():
        api_key = os.getenv('GROQ_API_KEY')
        if not api_key or not str(api_key).strip():
            raise ValueError(
                "GROQ_API_KEY environment variable is not set. "
                "Get a free key from https://console.groq.com/ (keys start with gsk_). "
                "Add it to .env in the project root or Backend folder."
            )
        if not str(api_key).strip().startswith('gsk_'):
            raise ValueError(
                "GROQ_API_KEY appears invalid (Groq keys start with gsk_). "
                "Get a free key from https://console.groq.com/ - xAI/OpenAI keys won't work."
            )
        client = _groq.Groq(api_key=api_key)
    return client

def get_required_libraries():
    global _PyPDF2, _Document, _pd, _CronTab, _np
    if _PyPDF2 is None:
        import PyPDF2
        _PyPDF2 = PyPDF2
    if _Document is None:
        from docx import Document
        _Document = Document
    if _pd is None:
        import pandas as pd
        _pd = pd
    if _CronTab is None:
        from crontab import CronTab
        _CronTab = CronTab
    if _np is None:
        import numpy as np
        _np = np
    return _PyPDF2, _Document, _pd, _CronTab, _np

def get_PyPDF2():
    return get_required_libraries()[0]

def get_Document():
    return get_required_libraries()[1]

def get_pd():
    return get_required_libraries()[2]

def get_CronTab():
    return get_required_libraries()[3]

def get_np():
    return get_required_libraries()[4]

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'csv'}

# Security configurations
MAX_FILE_SIZE = int(os.getenv('MAX_FILE_SIZE', 10485760))  # 10MB default
CORS_ORIGINS = os.getenv('CORS_ORIGINS', 'http://localhost:3000').split(',')
RATE_LIMIT_PER_MINUTE = int(os.getenv('RATE_LIMIT_PER_MINUTE', 10))
RATE_LIMIT_PER_HOUR = int(os.getenv('RATE_LIMIT_PER_HOUR', 100))

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

app = Flask(__name__)

# Security configurations
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', secrets.token_hex(32))
app.config['MAX_CONTENT_LENGTH'] = MAX_FILE_SIZE

# Configure CORS with specific origins
CORS(app, origins=CORS_ORIGINS, supports_credentials=True)

# Configure rate limiting
limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=[f"{RATE_LIMIT_PER_HOUR} per hour", f"{RATE_LIMIT_PER_MINUTE} per minute"]
)

def extract_text_from_file(file):
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower()
    
    try:
        if ext == 'pdf':
            PyPDF2_module = get_PyPDF2()
            pdf_reader = PyPDF2_module.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():  # Only add non-empty pages
                    text += page_text + "\n"
            return text.strip()
        elif ext == 'docx':
            Document_class = get_Document()
            doc = Document_class(file)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():  # Only add non-empty paragraphs
                    text += para.text + "\n"
            return text.strip()
        elif ext == 'csv':
            pd_module = get_pd()
            df = pd_module.read_csv(file)
            # Convert to a more readable format
            text = f"CSV Data with {len(df)} rows and {len(df.columns)} columns:\n\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            text += "First 10 rows:\n"
            text += df.head(10).to_string(index=False)
            return text
        elif ext == 'txt':
            file.seek(0)  # Reset file pointer
            content = file.read()
            if isinstance(content, bytes):
                content = content.decode('utf-8', errors='ignore')
            return content.strip()
        else:
            return f"Unsupported file type: {ext}"
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        return f"Error extracting text from file: {str(e)}"

# Constants
# Use raw string or forward slashes to avoid invalid escape sequences on Windows
CRON_COMMAND = r"Backend\subreddit_topics.json"  # Update with your actual cron script path
SUBREDDIT_JSON_PATH = os.path.join(os.path.dirname(__file__), "subreddit_topics.json")

# Load subreddit data
def load_subreddit_data():
    try:
        with open(SUBREDDIT_JSON_PATH, 'r') as file:
            return json.load(file)
    except Exception as e:
        print(f"Error loading subreddit_topics.json: {str(e)}")
        return {
            "TravelHacks": [
                "iterninary", "mexican", "searched", "prepared", 
                "pack", "lagging", "trolley"
            ],
            "CryptoCurrency": [
                "kanye", "adress", "brightly", "aped", 
                "pointless", "awakens", "tulsi"
            ]
        }

# Add this helper function at the top (or near your imports)
def convert_numpy_types(obj):
    if isinstance(obj, dict):
        return {k: convert_numpy_types(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_types(i) for i in obj]
    elif isinstance(obj, tuple):
        return tuple(convert_numpy_types(i) for i in obj)
    else:
        np_module = get_np()
        if isinstance(obj, np_module.generic):
            return obj.item()
        return obj

@app.route('/api/upload', methods=['POST'])
@limiter.limit("5 per minute")  # Rate limit file uploads
def upload_file():
    try:
        # Validate file presence
        if 'file' not in request.files:
            logger.warning("No file part in upload request")
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['file']
        if file.filename == '':
            logger.warning("No file selected for upload")
            return jsonify({'error': 'No selected file'}), 400
        
        # Validate file type
        if not allowed_file(file.filename):
            logger.warning(f"Invalid file type attempted: {file.filename}")
            return jsonify({'error': 'File type not allowed. Only PDF, DOCX, TXT, and CSV files are permitted.'}), 400
        
        # Validate file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > MAX_FILE_SIZE:
            logger.warning(f"File too large: {file_size} bytes")
            return jsonify({'error': f'File too large. Maximum size allowed is {MAX_FILE_SIZE // (1024*1024)}MB'}), 400
        
        # Extract text from the file
        extracted_text = extract_text_from_file(file)
        
        logger.info(f"Extracted text length: {len(extracted_text)} characters")
        logger.info(f"First 200 chars: {extracted_text[:200]}...")
        
        if not extracted_text or len(extracted_text.strip()) == 0:
            return jsonify({'error': 'No readable text found in the file'}), 400
        
        # Process with Groq
        client_instance = get_groq_client()
        chat_completion = client_instance.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a helpful assistant that analyzes documents and provides insights. Keep your responses concise and focused."},
                {"role": "user", "content": f"Please analyze this text and provide key insights: {extracted_text[:4000]}"}  # Limiting to 4000 chars for safety
            ],
            model="llama-3.1-8b-instant",
            temperature=0.7,
            max_tokens=1000,  # Limit response length
        )
        
        logger.info(f"Successfully processed file: {file.filename}")
        return jsonify({
            'filename': secure_filename(file.filename),
            'text_preview': extracted_text[:1000],  # Send preview of extracted text
            'analysis': chat_completion.choices[0].message.content
        })
    
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': 'An error occurred while processing the file'}), 500

# New: analyze multiple documents with optional user query, processed on Send
@app.route('/api/analyze-docs', methods=['POST'])
@limiter.limit("10 per minute")
def analyze_docs():
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files provided'}), 400

        files = request.files.getlist('files')
        user_query = request.form.get('query', '').strip()

        if not files or len(files) == 0:
            return jsonify({'error': 'No files provided'}), 400

        texts = []
        filenames = []
        total_chars = 0
        MAX_COMBINED_CHARS = 16000

        for f in files:
            if f and allowed_file(f.filename):
                txt = extract_text_from_file(f)
                if txt and txt.strip():
                    remaining = MAX_COMBINED_CHARS - total_chars
                    if remaining <= 0:
                        break
                    trimmed = txt[:remaining]
                    texts.append(trimmed)
                    filenames.append(secure_filename(f.filename))
                    total_chars += len(trimmed)

        if not texts:
            return jsonify({'error': 'No readable text could be extracted from the provided files'}), 400

        combined_text = "\n\n".join(texts)
        prompt = (
            (f"User query: {user_query}\n\n" if user_query else "") +
            "Analyze the following document text and provide concise, structured insights, key points, and any potential action items.\n\n" +
            combined_text
        )

        client_instance = get_groq_client()
        chat_completion = client_instance.chat.completions.create(
            messages=[
                {"role": "system", "content": "You are a helpful assistant that analyzes one or more documents and provides concise insights."},
                {"role": "user", "content": prompt[:4000]}
            ],
            model="llama-3.1-8b-instant",
            temperature=0.5,
            max_tokens=1000,
        )

        return jsonify({
            'filenames': filenames,
            'text_preview': combined_text[:1000],
            'analysis': chat_completion.choices[0].message.content
        })

    except Exception as e:
        logger.error(f"Error in analyze-docs: {str(e)}")
        traceback.print_exc()
        return jsonify({'error': 'An error occurred while analyzing documents'}), 500

@app.route('/api/chat', methods=['POST'])
@limiter.limit("20 per minute")  # Rate limit chat requests
def chat_endpoint():
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
            
        user_query = data.get('user_query')
        userID = data.get('userID')
        subreddit = data.get('subreddit')
        topics = data.get('topics')
        
        # Validate required parameters
        if not all([user_query, userID, subreddit, topics]):
            logger.warning(f"Missing required parameters in chat request: {data}")
            return jsonify({"error": "Missing required parameters: user_query, userID, subreddit, topics"}), 400
        
        # Validate input lengths
        if len(user_query) > 1000:
            return jsonify({"error": "Query too long. Maximum 1000 characters allowed."}), 400
        
        if not isinstance(topics, list) or len(topics) == 0:
            return jsonify({"error": "Topics must be a non-empty list"}), 400
        
        # Call kg_chat function with all topics
        kg_chat_module = import_kg_chat()
        # Ensure cache is initialized (idempotent, safe to call multiple times)
        if hasattr(kg_chat_module, 'initialize_kg_cache'):
            kg_chat_module.initialize_kg_cache()
        response = kg_chat_module.chat_with_kg(user_query, userID, subreddit, topics)
        logger.info(f"Successfully processed chat request for user: {userID}")
        return jsonify({"response": response})
    except Exception as e:
        logger.error(f"Error in chat endpoint: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": "An error occurred while processing your request"}), 500

@app.route('/api/visualize', methods=['POST'])
@limiter.limit("30 per minute")  # Rate limit visualization requests
def visualize_endpoint():
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400
            
        user_query = data.get('user_query')
        response = data.get('response')
        
        # Handle response if it's a dictionary
        if isinstance(response, dict) and 'response' in response:
            response = response['response']
        
        if not all([user_query, response]):
            logger.warning("Missing user_query or response in visualize request")
            return jsonify({"error": "Missing user_query or response"}), 400
        
        # Validate input lengths
        if len(user_query) > 1000:
            return jsonify({"error": "Query too long. Maximum 1000 characters allowed."}), 400
        
        if len(str(response)) > 5000:
            return jsonify({"error": "Response too long. Maximum 5000 characters allowed."}), 400
        
        # Get visualization recommendations
        vrs_module = import_vrs()
        recommended_charts = vrs_module.getViz(user_query, response)
        logger.info("Successfully generated visualization recommendations")
        return jsonify(convert_numpy_types(recommended_charts))
    except Exception as e:
        logger.error(f"Error in visualize endpoint: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": "An error occurred while generating visualizations"}), 500

@app.route('/api/subreddits', methods=['GET'])
def subreddits_endpoint():
    # Load subreddits and topics from JSON file
    subreddit_topics = load_subreddit_data()
    return jsonify(subreddit_topics)

@app.route('/api/schedule', methods=['POST'])
def schedule_endpoint():
    data = request.json
    cronInterval = int(data.get('cronInterval', 0))
    cronStart = int(data.get('cronTime', 0))
    userID = data.get('userID')
    subreddit = data.get('subreddit')
    topics = data.get('topics')

    print(f"cron start: {cronStart}")
    print(f"cronInterval: {cronInterval}")
    print(f"Setting up cron for user: {userID}, subreddit: {subreddit}")

    try:
        if not cronStart and not cronInterval:
            print("empty")
            return jsonify({"error": "Nothing selected"}), 400

        if cronStart < 1 or cronStart > 23:
            print("start time not in range")
            return jsonify({"error": "invalid start time"})

        CronTab_module = get_CronTab()
        cron = CronTab_module(user=True)

        for job in cron:
            if job.command == CRON_COMMAND:
                print("cron exists")
                cron.remove(job)
                print("cron removed")

        new_job = cron.new(command=CRON_COMMAND)
        print(new_job, "\n")

        new_job.setall(f"{cronStart} */{cronInterval} * * *")
        print("cron set")
        cron.write()

        print("cron applied")
        return jsonify({"message": f"cronjob successfully applied at {cronStart} every {cronInterval} minutes"}), 200

    except Exception as e:
        print(f"Error setting up cron job: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": str(e)})

@app.route('/api/metrics', methods=['GET'])
def metrics():
    """Return process and host CPU/RAM metrics for deployment sizing."""
    try:
        import psutil
        proc = psutil.Process()
        mem = proc.memory_info()
        cpu_percent = proc.cpu_percent(interval=0.1)
        payload = {
            "process": {
                "cpu_percent": round(cpu_percent, 2),
                "memory_rss_mb": round(mem.rss / (1024 * 1024), 2),
                "memory_vms_mb": round(mem.vms / (1024 * 1024), 2),
            },
        }
        try:
            vm = psutil.virtual_memory()
            payload["host"] = {
                "memory_total_mb": round(vm.total / (1024 * 1024), 2),
                "memory_available_mb": round(vm.available / (1024 * 1024), 2),
                "cpu_count_logical": psutil.cpu_count(),
            }
        except Exception:
            payload["host"] = None
        return jsonify(payload)
    except Exception as e:
        logger.exception("Metrics collection failed")
        return jsonify({"error": "Metrics unavailable", "detail": str(e)}), 500


@app.route('/', methods=['GET'])
def health_check():
    return jsonify({"status": "API is running"})

if __name__ == '__main__':
    app.run(debug=True, port=5000)